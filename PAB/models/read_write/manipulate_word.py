from docx import Document
from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Cm, Pt

def create_word_table(doc, headers, data):
    # добавляем таблицу с одной строкой для заполнения названий колонок
    table = doc.add_table(1, len(headers)+1)
    # определяем стиль таблицы
    table.style = 'Light Shading Accent 1'
    table.autofit = True
    # Получаем указатель на строку с названием колонок из добавленной таблицы
    head_row = table.rows[0].cells
    # добавляем названия колонок
    for idx, name in enumerate(headers):
        head_row[idx].width = Mm(150)
        paragraf = head_row[idx].paragraphs[0]
        # название колонки
        paragraf.add_run(name).bold = True
        # выравниваем посередине
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # добавляем данные к существующей таблице
    for row in data:
        # добавляем строку с ячейками к объекту таблицы
        add_row = table.add_row().cells
        for idx, item in enumerate(row):
            # вставляем данные в ячейки
            paragraf = add_row[idx].paragraphs[0]
            paragraf.add_run(str(item)).bold = False
            paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_myword(sets, file_name):
    #Создаем документ
    doc = Document()
    # доступ к первой секции:
    section = doc.sections[0]
    # высота листа в сантиметрах
    section.page_height = Cm(29.7)
    # ширина листа в сантиметрах
    section.page_width = Cm(21.0)
    # левое поле в миллиметрах
    section.left_margin = Mm(20.4)
    # правое поле в миллиметрах
    section.right_margin = Mm(10)
    # изменяем стиль текста по умолчанию
    style = doc.styles['Normal']
    # название шрифта по умолчанию
    style.font.name = 'Times New Roman'
    # размер шрифта по умолчанию
    style.font.size = Pt(12)
    for idx, set in enumerate(sets):
        paragraf = doc.add_paragraph(f'Таблица № {idx + 1} "{set.data_lable}"')
        paragraf_format = paragraf.paragraph_format
        paragraf_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraf_format.left_indent = Mm(0)
        paragraf_format.space_before = Pt(20)
        paragraf_format.space_after = Pt(10)
        create_word_table(doc, set.data_header, set.data_table)
    try:
        doc.save(file_name)
    except Exception as other:
        Error_str = f'Не удалось сохранить файл. Ошибка {other}'
        print(Error_str)

def read_myword(file_name):
    pass
