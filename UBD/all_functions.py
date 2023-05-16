import pandas as pd
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm, Cm, Pt
from pathlib import Path

import app_logger
logger = app_logger.get_logger(__name__)

def read_data(file_name, format, about, **kwargs):
    read_com = {"csv": pd.read_csv, "xlsx": pd.read_excel}
    data_path = Path(file_name).absolute()
    try:
        data = read_com[format](data_path, **kwargs)
    except Exception as err:
        print(f"При считывании данных c {about} возникла ошибка {err}")
        file_name = input(f"Введите путь до файла с {about}: ")
        data = read_data(file_name, format, about, **kwargs)
    return data

"""
Семейство функций для обработки ДФ
"""
#Функция для выявления из списка наиболее критического уровня опасности для каждой уязвимости
def calc_max_level(one_str):
    rang_levels = {'Критический': 1, 'Высокий': 2, 'Средний': 3, 'Низкий':4, 'Нет': 5}
    rang_levels_rev = {1: 'Критический', 2: 'Высокий', 3: 'Средний', 4: 'Низкий', 5: 'Нет', 6: 'Ошибка'}
    levels = []
    for one in one_str.split('\n'):
        level = rang_levels.get(one.split(' ')[0], 0)
        if not level:
            Error = f"Не удалось подобрать уровень критичностию. Метод calc_max_level, строка: {one}"
            logger.warning(Error)
        else:
            levels.append(level)
    try:
        new_level = rang_levels_rev[min(levels)]
    except:
        new_level = 'Ошибка'
    return new_level


#Функция для преобразования столбца с данными в тип datetime с обработкой невернозаписанных дат
def change_date(str_d):
    try:
        date = pd.to_datetime(str_d, format='%d.%m.%Y')
    except:
        str_d = input(f'В БД выявлена некорректная дата: {str_d}, введите корректную дату в формате ДД.ММ.ГГГГ: ')
        date = change_date(str_d)
    return date

"""
Семейство функций для расчета различных вспомогательных словарей сопоставления 
и отбора данных, релевантных установленному ПО
"""
#Функция для расчета словаря-сопоставления {название версии: название ПО}
def calc_bd_version_dict(data):
    logger.info("Расчет словаря - сопоставления {название версии: название ПО}. Метод calc_bd_version_dict ...")
    #РАсчет всех названия версий
    all_versions = list(set(', '.join(data.version).split(',')))

    #Расчет словаря-сопоставления {название версии: название ПО}
    bd_vers_names = {}
    template = re.compile(r'(.+)([(]{1})(.+)([)]{1})')
    for version in all_versions:
        try:
            assign_name = re.search(template, version)[3]
            bd_vers_names[version] = assign_name
        except:
            bd_vers_names[version] = version
    return bd_vers_names


#Функция для фильтрации БДУ по заданному перечню ПО
def search_common_soft(one_str, bd_pc_soft):
    if type(one_str) == str:
        ours = []
        for one_element in one_str.split(','):
            if bd_pc_soft.get(one_element):
                ours.append(one_element)
            else:
                pass
        if len(ours) > 0:
            return ', '.join(ours)
        else:
            return 'no'
    else:
        logger.warning(f'Не удалось подобрать подходящее ПО в строке {one_str}')
    return "one_str"


#Функция для расчета словаря - соответсвия {название из бду: [названия из установленного ПО]}
def search_common_dict(soft_bd, soft_pc, k1, k2):
    logger.info("Расчет словаря - соответсвия {название из бду: [названия из установленного ПО]. Метод search_common_dict ...")
    ours = {}
    for one_bd in soft_bd:
        for one_pc in soft_pc:
            if calc_equal(norm_sentence(one_bd), norm_sentence(one_pc), 3, k1) > k2:
                if ours.get(one_bd):
                    ours[one_bd].append(one_pc)
                else:
                     ours[one_bd] = [one_pc]
            else:
                pass
    ours = {key: list(set(value)) for key, value in ours.items()}
    return ours


"""
Семейство функций для подготовки отчета в формате word
"""

#Функция для создания таблицы в word
#Функция для создания таблицы в word
def create_word_table(doc, data, dict_indexes, widths):
    # добавляем таблицу с одной строкой для заполнения названий колонок
    table = doc.add_table(1, len(data.columns))
    # определяем стиль таблицы
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    # Получаем указатель на строку с названием колонок из добавленной таблицы
    head_row = table.rows[0].cells
    # добавляем названия колонок
    for idx, name in enumerate(data.columns):
        head_row[idx].width = Mm(widths[idx])
        paragraf = head_row[idx].paragraphs[0]
        # название колонки
        paragraf.add_run(name).bold = True
        # выравниваем по ширине посередине
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #Расчитываем длину словаря с объединенными стороками
    # добавляем данные к существующей таблице
    for ind_row in range((data.shape[0])):
      if dict_indexes:
        name = dict_indexes.get(ind_row, None)
      else:
        name = None
      # добавляем строку с ячейками к объекту таблицы
      if name:
        add_row = table.add_row().cells
        a = add_row[0]
        b = add_row[data.shape[-1] - 1]
        A = a.merge(b)
        paragraf = A.paragraphs[0]
        paragraf.add_run(name).bold = True
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        A.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
      add_row = table.add_row().cells
      for ind_col in range(data.shape[-1]):
          # вставляем данные в ячейки
          paragraf = add_row[ind_col].paragraphs[0]
          paragraf.add_run(str(data.values[ind_row, ind_col])).bold = False
          paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
          #Выравнивание в ячейке по центру
          add_row[ind_col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

#Функция для создания документа Word
def create_myword(sets: list, file_name, dict_indexes: list, widths: list):
    #Создаем документ
    doc = Document()
    # доступ к первой секции:
    section = doc.sections[0]
    # ориетация страницы
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # левое поле в миллиметрах
    section.left_margin = Mm(10)
    # правое поле в миллиметрах
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(5)
    # изменяем стиль текста по умолчанию
    style = doc.styles['Normal']
    # название шрифта по умолчанию
    style.font.name = 'Times New Roman'
    # размер шрифта по умолчанию
    style.font.size = Pt(9)
    for idx, data in enumerate(sets):
        paragraf = doc.add_paragraph(f'Таблица № {idx + 1}')
        paragraf_format = paragraf.paragraph_format
        paragraf_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraf_format.left_indent = Mm(0)
        paragraf_format.space_before = Pt(20)
        paragraf_format.space_after = Pt(10)
        create_word_table(doc, data, dict_indexes[idx], widths[idx])
    try:
        doc.save('C:\\fstek_ready\\'+file_name)
    except Exception as other:
        logger.error(f'Не удалось сохранить файл. Ошибка {other}')

"""
Семейство функций, чтобы сопоставить название ПО из ПК и ПО из БД 
(подобрать коррелированные друг с другом строчки)
"""

#Отчищает строку от мусора
def norm_sentence(soft_name):
    #Строчные буквы
    soft_name = soft_name.lower()
    #Исключаем кавычки, переводы строки
    soft_name = re.sub("""^\s+|'|"|\n|\r|\s+$""", '', soft_name)
    #Исключаем повторяющиеся пробелы
    soft_name = " ".join(soft_name.split())
    return soft_name

#Разделяет строку на кусочкис наложением
def get_tokens(soft_name):
    words = soft_name.split(' ')
    tokens = [word for word in words if len(word) > 2]
    return tokens

#Проверяет похожесть двух кусочков
def is_tokens_equal(first_token, second_token, subtoken_len, true_limit):
    first_number = len(first_token) - subtoken_len + 1
    second_number = len(second_token) - subtoken_len + 1
    used_tokens = [False for ind in range(second_number)]
    equal_count = 0
    for first_ind in range(first_number):
        subtoken_first = first_token[first_ind : first_ind + subtoken_len]
        for second_ind in range(second_number):
            if not used_tokens[second_ind]:
                subtoken_second = second_token[first_ind : first_ind + subtoken_len]
                if subtoken_first == subtoken_second:
                    equal_count = equal_count + 1
                    used_tokens[second_ind] = True
                    break
    subtoken_first_count = len(first_token) - subtoken_len + 1
    subtoken_second_count = len(second_token) - subtoken_len + 1
    tanimoto = (1.0 * equal_count) / (subtoken_first_count + subtoken_second_count - equal_count)
    return tanimoto > true_limit

#Перевирает все кусочки из двух строк и собирает только одинаковые в отдельный список
def get_equals_tokens(tokens_first, tokens_second, subtoken_len, true_limit):
    equals_tokens = []
    used_token = [False for ind in range(len(tokens_second))]
    for first_ind in range(len(tokens_first)):
        for second_ind in range(len(tokens_second)):
            if not used_token[second_ind]:
                if is_tokens_equal(tokens_first[first_ind], tokens_second[second_ind], subtoken_len, true_limit):
                    equals_tokens.append(tokens_first[first_ind])
                    used_token[second_ind] = True
                    break
    return equals_tokens

#Определяет степень одинаковости строк
def calc_equal(first_word, second_word, subtoken_len, true_limit):
    tokens_first = get_tokens(norm_sentence(first_word))
    tokens_second = get_tokens(norm_sentence(second_word))
    equals_tokens = get_equals_tokens(tokens_first, tokens_second, subtoken_len, true_limit)
    equals_count = len(equals_tokens)
    first_count = len(tokens_first)
    second_count = len(tokens_second)
    if first_count + second_count - equals_count != 0:
      result_value = (1.0 * equals_count) / (first_count + second_count - equals_count)
    else:
      result_value = 1
    return result_value

"""
Семейство функций чтобы выявить версии ПО из БД, 
которые пересекаются с версиями ПО из ПК
"""
#Функция, которая на основании словаря {версии БД: названия ПО БД} и
#словаря {названия ПО БД: названия ПО ПК} определяет список версий из БД,
#актуальных для ПО из ПК
def calc_actual_versions(bd_vers_names, bd_pc_soft, name_vers_pc):
    logger.info("Расчет списка версий из БД, актуальных для ПО из ПК. Метод calc_actual_versions ... ")
    actuals = []
    for vers_from_bd in bd_vers_names.keys():
        actual = False
        #Ищем устновленное ПО, связанное с ПО из БД
        name_from_bd = bd_vers_names[vers_from_bd]
        list_name_from_pc = bd_pc_soft[name_from_bd]
        #Для каждого установленного ПО ищем все его версии
        for name_from_pc in list_name_from_pc:
            if actual == True:
                break
            else:
                #Для каждой версии установленного По проверяем актуальность
                actual = check_actual(name_vers_pc, name_from_pc, vers_from_bd)
        if actual == True:
            actuals.append(vers_from_bd)
        else:
            pass
    return actuals

#Проверяет актальность версий установленного ПО
def check_actual(name_vers_pc, name_from_pc, vers_from_bd):
    actual = False
    vers_from_pc = name_vers_pc[name_from_pc]
    if vers_from_pc:
        if type(vers_from_pc) != list:
            if str(vers_from_pc) == 'nan':
                actual == False
            else:
                actual = decision_by_vers(vers_from_bd, vers_from_pc)
        else:
            for ver_from_pc in vers_from_pc:
                if actual == True:
                    break
                elif str(ver_from_pc) == 'nan':
                    actual == False
                else:
                    actual = decision_by_vers(vers_from_bd, ver_from_pc)
    return actual

#Определяет атуальность для одной записи
def decision_by_vers(vers_string, pc_string):
    template_eq = re.compile(r'(\b)((\d{1,}[.]{,1})+)(\b)')
    if len(pc_string) < 2:
        return False
    pc_numbers = calc_comb_numb(pc_string, template_eq, 2)
    t, ot, do, eq = calc_vers_type(vers_string)
    if t == "do":
        actual = decision_do(do, pc_numbers)
    elif t == "ot":
        actual = decision_ot(ot, pc_numbers)
    elif t == "between":
        actual_do = decision_do(do, pc_numbers)
        actual_ot = decision_ot(ot, pc_numbers)
        if actual_ot and actual_do:
            return True
        else:
            return False
    else:
        if eq == pc_numbers:
            return True
        else:
            return False
    return actual

#Определяет актульность в случае,если у нас версия записана, как до <последовательность цифр>
def decision_do(bd_numbers, pc_numbers):
    actual = True
    max_ind = min([len(bd_numbers), len(pc_numbers)])
    for ind in range(max_ind):
        if pc_numbers[ind] < bd_numbers[ind]:
            #Считаем что ubd - актуальна
            actual = True
            break
        elif pc_numbers[ind] == bd_numbers[ind]:
            #Проверяем следующую позицию
             pass
        else:
            #Для данного ПО с ПК уязвимость неактуальна, переодим к следующему
            actual = False
            break
    return actual

#Определяет актульность в случае,если у нас версия записана, как от <последовательность цифр>
def decision_ot(bd_numbers, pc_numbers):
    actual = True
    max_ind = min([len(bd_numbers), len(pc_numbers)])
    for ind in range(max_ind):
        if pc_numbers[ind] > bd_numbers[ind]:
            #Считаем что ubd - актуальна
            actual = True
            break
        elif pc_numbers[ind] == bd_numbers[ind]:
            #Проверяем следующую позицию
            pass
        else:
            #Для данного ПО с ПК уязвимость неактуальна, переодим к следующему
            actual = False
            break
    return actual


#Функция которая определяет как задана версия в столбце версий, варианты:
#от <версия>
#до <версия>
#от <версия> до <версия>
#<версия>
#{номер комбинации цифр: список цифр из комбинации}
def calc_vers_type(vers_string):
    vers = vers_string.split()
    template_do = re.compile(r'(.+)(до\s+)(\b)((\d{1,}[.]{,1})+)(\b)')
    template_ot = re.compile(r'(.+)(от\s+)(\b)((\d{1,}[.]{,1})+)(\b)')
    template_eq = re.compile(r'(\b)((\d{1,}[.]{,1})+)(\b)')
    if "от" in vers and "до" in vers:
        t = "between"
        ot = calc_comb_numb(vers_string, template_ot, 4, "от")
        do = calc_comb_numb(vers_string, template_do, 4, "до")
        eq = None
    elif "от" not in vers and "до" in vers:
        t, ot, do, eq = "do", None, calc_comb_numb(vers_string, template_do, 4, "до"), None
    elif "от" in vers and "до" not in vers:
        t, ot, do, eq = "ot", calc_comb_numb(vers_string, template_ot, 4, "от"), None, None
    else:
        t = "eq"
        eq = calc_comb_numb(vers_string, template_eq, 2, "равно")
        ot = None
        do = None
    return (t, ot, do, eq)

#Функция которая формирует список из найденных по шаблону цифр версии:
def calc_comb_numb(vers_string, template, ind_template, pref = None):
    #Любая комбинация цифр и точек
    match_list = re.search(template, vers_string)
    if match_list:
        bd_numbers = [int(pos) for pos in match_list[ind_template].split('.')]
    else:
        template = re.compile(r'(\b)((\d{1,}[.]{,1})+)(\b)')
        vers_string = input(f'''Для строки {vers_string} не удалось выделить данные версии с префиксом {pref}. \n
        Введите версию вручную: ''')
        bd_numbers = calc_comb_numb(vers_string, template, 2)
    return bd_numbers
